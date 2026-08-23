"""DOCX/PDF/PPT Vision 处理的公开契约回归。"""

from __future__ import annotations

import asyncio
import multiprocessing
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from PIL import Image

from fast_app.core.config import Settings
from fast_app.domain.knowledge_models import (
    LoadedWordDocument,
    VisionAnalysisResult,
    VisionImageContent,
    VisionImageOccurrence,
    WordBlock,
)
from fast_app.ingestion.processing.document_vision import DocumentVisionService
from fast_app.ingestion.processing.word_processing import WordDocumentLoader
from fast_app.ingestion.processing.pdf_processing import PdfDocumentLoader
from fast_app.ingestion.processing.chunk_builders import ChunkBuildOptions
from fast_app.ingestion.processing.office_chunk_builders import WordChunkBuilder
from fast_app.ingestion.processing.structured_document_processor import (
    DocumentProcessingError,
    StructuredDocumentProcessor,
)
from fast_app.ingestion.validation.document_validation import (
    DocumentPackageValidationError,
    validate_knowledge_document_package,
)


class RecordingVisionClient:
    """按 content identity 记录调用次数的确定性 Vision Adapter。"""

    def __init__(self) -> None:
        self.calls: list[str] = []

    async def analyze(self, *, content, mode, before_provider_call=None):
        if before_provider_call is not None:
            await before_provider_call()
        self.calls.append(content.content_id)
        return VisionAnalysisResult(
            extracted_text="diagram text",
            summary="architecture diagram",
            table_markdown=None,
            visual_facts=["Milvus precedes rerank"],
        )


class SelectiveFailureVisionClient(RecordingVisionClient):
    def __init__(self, failing_content_id: str) -> None:
        super().__init__()
        self.failing_content_id = failing_content_id

    async def analyze(self, *, content, mode, before_provider_call=None):
        if before_provider_call is not None:
            await before_provider_call()
        self.calls.append(content.content_id)
        if content.content_id == self.failing_content_id:
            from fast_app.components.vision.base import VisionAnalysisError

            raise VisionAnalysisError("VISION_PROVIDER_FAILED", "图片模型调用失败")
        return VisionAnalysisResult(
            extracted_text="ok",
            summary="usable image",
            table_markdown=None,
            visual_facts=[],
        )


def png_bytes() -> bytes:
    stream = BytesIO()
    Image.new("RGB", (16, 16), color="red").save(stream, format="PNG")
    return stream.getvalue()


def _write_shared_cache(cache_dir: str, summary: str) -> None:
    settings = Settings(
        _env_file=None,
        VISION_CACHE_ENABLED=True,
        VISION_CACHE_DIR=cache_dir,
    )
    service = DocumentVisionService(settings=settings)
    content = VisionImageContent.from_raw(png_bytes(), media_type="image/png")
    service._write_cache(
        content,
        "embedded_image",
        VisionAnalysisResult(
            extracted_text="",
            summary=summary,
            table_markdown=None,
            visual_facts=[],
        ),
    )


def test_document_validation_dispatches_docx_and_pdf(root: Path) -> None:
    from docx import Document
    from pypdf import PdfWriter

    docx_path = root / "sample.docx"
    document = Document()
    document.add_paragraph("hello")
    document.save(docx_path)
    assert validate_knowledge_document_package(
        docx_path, document_type="word"
    ).document_type == "word"

    pdf_path = root / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_path.open("wb") as stream:
        writer.write(stream)
    assert validate_knowledge_document_package(
        pdf_path, document_type="pdf"
    ).page_count == 1

    try:
        validate_knowledge_document_package(pdf_path, document_type="word")
    except DocumentPackageValidationError as exc:
        assert exc.code in {"INVALID_OOXML_PACKAGE", "DOCUMENT_TYPE_MISMATCH"}
    else:
        raise AssertionError("PDF 不能通过 DOCX OOXML 校验")


async def test_same_content_keeps_occurrences_but_reuses_analysis() -> None:
    raw = png_bytes()
    content = VisionImageContent.from_raw(raw, media_type="image/png")
    occurrences = [
        VisionImageOccurrence(
            occurrence_id=f"imgocc:docx:doc:block-{index}:rId1:1:{content.content_id}",
            content_id=content.content_id,
            source_locator=f"body/paragraph[{index}]",
            block_id=f"block-{index}",
            relationship_id="rId1",
            occurrence_index=index,
        )
        for index in (1, 2)
    ]
    client = RecordingVisionClient()
    service = DocumentVisionService(
        settings=Settings(_env_file=None, VISION_ENABLED=True),
        client=client,
    )
    results = await service.analyze_assets(
        contents={content.content_id: content},
        occurrences=occurrences,
        mode="embedded_image",
    )
    assert client.calls == [content.content_id]
    assert set(results) == {item.occurrence_id for item in occurrences}
    assert results[occurrences[0].occurrence_id] == results[occurrences[1].occurrence_id]


async def test_disabled_vision_never_constructs_or_calls_client() -> None:
    raw = png_bytes()
    content = VisionImageContent.from_raw(raw, media_type="image/png")
    occurrence = VisionImageOccurrence(
        occurrence_id=f"imgocc:ppt:doc:1:2:{content.content_id}",
        content_id=content.content_id,
        source_locator="slide[1]/shape[2]",
        page_or_slide_number=1,
        anchor_id="2",
        occurrence_index=1,
    )
    service = DocumentVisionService(
        settings=Settings(_env_file=None, VISION_ENABLED=False),
        client_factory=lambda: (_ for _ in ()).throw(
            AssertionError("disabled Vision must not construct a client")
        ),
    )
    outcome = await service.analyze_assets_with_warnings(
        contents={content.content_id: content},
        occurrences=[occurrence],
        mode="embedded_image",
    )
    assert outcome.results == {}
    assert [warning.code for warning in outcome.warnings] == ["VISION_DISABLED"]


async def test_one_image_failure_does_not_discard_other_image_results() -> None:
    first = VisionImageContent.from_raw(png_bytes(), media_type="image/png")
    blue = BytesIO()
    Image.new("RGB", (16, 16), color="blue").save(blue, format="PNG")
    second = VisionImageContent.from_raw(blue.getvalue(), media_type="image/png")
    occurrences = [
        VisionImageOccurrence(
            occurrence_id=f"occ:{index}",
            content_id=content.content_id,
            source_locator=f"slide[1]/shape[{index}]",
            page_or_slide_number=1,
            occurrence_index=index,
        )
        for index, content in enumerate((first, second), start=1)
    ]
    settings = Settings(_env_file=None, VISION_ENABLED=True)
    client = SelectiveFailureVisionClient(first.content_id)
    outcome = await DocumentVisionService(settings=settings, client=client).analyze_assets_with_warnings(
        contents={first.content_id: first, second.content_id: second},
        occurrences=occurrences,
        mode="embedded_image",
    )
    assert set(client.calls) == {first.content_id, second.content_id}
    assert set(outcome.results) == {occurrences[1].occurrence_id}
    assert [(item.code, item.source_locator) for item in outcome.warnings] == [
        ("VISION_PROVIDER_FAILED", occurrences[0].source_locator)
    ]


async def test_lease_loss_prevents_next_provider_call() -> None:
    first = VisionImageContent.from_raw(png_bytes(), media_type="image/png")
    blue = BytesIO()
    Image.new("RGB", (16, 16), color="blue").save(blue, format="PNG")
    second = VisionImageContent.from_raw(blue.getvalue(), media_type="image/png")
    occurrences = [
        VisionImageOccurrence(
            occurrence_id=f"lease-occ:{index}",
            content_id=content.content_id,
            source_locator=f"paragraph[{index}]",
            occurrence_index=index,
        )
        for index, content in enumerate((first, second), start=1)
    ]
    checks = 0

    async def assert_owned() -> None:
        nonlocal checks
        checks += 1
        if checks > 1:
            raise RuntimeError("lease lost")

    client = RecordingVisionClient()
    settings = Settings(
        _env_file=None, VISION_ENABLED=True, VISION_MAX_CONCURRENCY=1
    )
    try:
        await DocumentVisionService(
            settings=settings, client=client
        ).analyze_assets_with_warnings(
            contents={first.content_id: first, second.content_id: second},
            occurrences=occurrences,
            mode="embedded_image",
            before_external_call=assert_owned,
        )
    except RuntimeError as exc:
        assert str(exc) == "lease lost"
    else:
        raise AssertionError("lease lost 必须中止整个文档处理，而不是降级为图片 warning")
    assert len(client.calls) == 1


def test_multi_process_cache_publish_is_atomic(root: Path) -> None:
    cache_dir = root / "vision-cache"
    context = multiprocessing.get_context("spawn")
    processes = [
        context.Process(target=_write_shared_cache, args=(str(cache_dir), summary))
        for summary in ("worker-a", "worker-b")
    ]
    for process in processes:
        process.start()
    for process in processes:
        process.join(20)
        assert process.exitcode == 0
    files = list(cache_dir.glob("*.json"))
    assert len(files) == 1
    value = VisionAnalysisResult.model_validate_json(files[0].read_text(encoding="utf-8"))
    assert value.summary in {"worker-a", "worker-b"}
    assert not list(cache_dir.glob("*.tmp"))


def _build_docx_with_repeated_image(path: Path, image_path: Path, *, prefix: bool) -> None:
    from docx import Document

    document = Document()
    if prefix:
        document.add_paragraph("unrelated preface")
    document.add_heading("Stable Section", level=1)
    first = document.add_paragraph("first block")
    first.add_run().add_picture(str(image_path))
    second = document.add_paragraph("second block")
    second.add_run().add_picture(str(image_path))
    document.save(path)


def test_word_occurrence_and_block_identity(root: Path) -> None:
    image_path = root / "same.png"
    image_path.write_bytes(png_bytes())
    first_path = root / "first.docx"
    second_path = root / "second.docx"
    _build_docx_with_repeated_image(first_path, image_path, prefix=False)
    _build_docx_with_repeated_image(second_path, image_path, prefix=True)

    loader = WordDocumentLoader()
    first = loader.load_structured_file(first_path, source_path="docs/sample.docx")
    repeated = loader.load_structured_file(first_path, source_path="docs/sample.docx")
    shifted = loader.load_structured_file(second_path, source_path="docs/sample.docx")

    assert len(first.vision_contents) == 1
    assert len(first.vision_occurrences) == 2
    assert first.vision_occurrences[0].occurrence_id != first.vision_occurrences[1].occurrence_id
    assert [item.occurrence_id for item in first.vision_occurrences] == [
        item.occurrence_id for item in repeated.vision_occurrences
    ]

    first_section_blocks = [block.block_id for block in first.blocks if block.section_title == "Stable Section"]
    shifted_section_blocks = [block.block_id for block in shifted.blocks if block.section_title == "Stable Section"]
    assert first_section_blocks == shifted_section_blocks


def test_pdf_zero_text_page_renders_full_page(root: Path) -> None:
    from pypdf import PdfWriter

    path = root / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=100)
    with path.open("wb") as stream:
        writer.write(stream)
    settings = Settings(
        _env_file=None,
        PDF_RENDER_SCALE=2.0,
        PDF_MIN_RENDER_SCALE=0.5,
        VISION_MAX_IMAGE_PIXELS=1_000_000,
    )
    document = PdfDocumentLoader(settings).load_structured_file(
        path, source_path="docs/blank.pdf"
    )
    assert len(document.pages) == 1
    assert document.pages[0].scanned_candidate is True
    assert document.pages[0].native_text == ""
    assert len(document.pages[0].vision_occurrence_ids) == 1
    occurrence = document.vision_occurrences[0]
    assert occurrence.occurrence_id.endswith(":page:1:scanned")
    assert document.vision_contents[occurrence.content_id].width == 400
    assert document.vision_contents[occurrence.content_id].height == 200


def test_pdf_ignores_inherited_images_not_drawn_on_page(root: Path) -> None:
    path = root / "resource-table.pdf"
    path.write_bytes(b"not-read-by-fake")
    page = SimpleNamespace(
        extract_text=lambda: "This page has enough native text to stay native.",
        images=[
            SimpleNamespace(
                data=png_bytes(), name="unused.png", is_displayed=False
            ),
            SimpleNamespace(
                data=png_bytes(), name="drawn.png", is_displayed=True
            ),
        ],
    )
    with patch(
        "fast_app.ingestion.processing.pdf_processing.PdfReader",
        return_value=SimpleNamespace(pages=[page]),
    ):
        document = PdfDocumentLoader(
            Settings(_env_file=None, PDF_SCANNED_TEXT_THRESHOLD=1)
        ).load_structured_file(path, source_path="docs/resource-table.pdf")
    assert len(document.vision_occurrences) == 1
    assert document.vision_occurrences[0].anchor_id == "drawn.png"


async def test_word_vision_text_is_bound_to_its_block(root: Path) -> None:
    image_path = root / "bound.png"
    image_path.write_bytes(png_bytes())
    path = root / "bound.docx"
    _build_docx_with_repeated_image(path, image_path, prefix=False)
    settings = Settings(_env_file=None, VISION_ENABLED=True)
    client = RecordingVisionClient()
    processor = StructuredDocumentProcessor(
        settings=settings,
        vision_service=DocumentVisionService(settings=settings, client=client),
    )
    result = await processor.process_file(
        path,
        document_type="word",
        source_path="docs/bound.docx",
        options=ChunkBuildOptions(
            source="test", max_chars=400, overlap_chars=0, max_tokens=400, min_chars=1
        ),
    )
    assert result.parents == []
    assert len(result.chunks) >= 1
    assert sum(chunk.content.count("architecture diagram") for chunk in result.chunks) >= 2
    assert all("data:image" not in chunk.content for chunk in result.chunks)
    assert all("normalized_bytes" not in chunk.metadata for chunk in result.chunks)


async def test_scanned_pdf_without_vision_is_fatal(root: Path) -> None:
    from pypdf import PdfWriter

    path = root / "disabled-scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as stream:
        writer.write(stream)
    settings = Settings(_env_file=None, VISION_ENABLED=False)
    processor = StructuredDocumentProcessor(settings=settings)
    try:
        await processor.process_file(
            path,
            document_type="pdf",
            source_path="docs/disabled-scan.pdf",
            options=ChunkBuildOptions(
                source="test", max_chars=400, overlap_chars=0, max_tokens=400, min_chars=1
            ),
        )
    except DocumentProcessingError as exc:
        assert exc.code == "PDF_PAGE_CONTENT_UNAVAILABLE"
    else:
        raise AssertionError("扫描页无正文且 Vision 关闭时必须阻止导入")


def test_oversized_word_block_has_unique_stable_chunk_ids() -> None:
    block = WordBlock(
        block_id="paragraph-stable",
        block_type="paragraph",
        text="alpha beta gamma delta " * 30,
        section_id="section-stable",
        section_title="Stable Section",
        heading_level=1,
    )
    document = LoadedWordDocument(
        source_path="docs/long.docx",
        blocks=[block],
        metadata={"doc_id": "doc-long", "document_type": "word"},
    )
    chunks = WordChunkBuilder().build(
        document,
        ChunkBuildOptions(
            source="test", max_chars=90, overlap_chars=10, max_tokens=90, min_chars=1
        ),
        embedding_fingerprint="embedding-v1",
    )
    assert len(chunks) > 1
    assert len({chunk.id for chunk in chunks}) == len(chunks)
    assert len({chunk.metadata["identity_key"] for chunk in chunks}) == len(chunks)


async def main() -> None:
    with TemporaryDirectory() as directory:
        root = Path(directory)
        test_document_validation_dispatches_docx_and_pdf(root)
        test_word_occurrence_and_block_identity(root)
        test_pdf_zero_text_page_renders_full_page(root)
        test_pdf_ignores_inherited_images_not_drawn_on_page(root)
        await test_word_vision_text_is_bound_to_its_block(root)
        await test_scanned_pdf_without_vision_is_fatal(root)
        test_oversized_word_block_has_unique_stable_chunk_ids()
        test_multi_process_cache_publish_is_atomic(root)
    await test_same_content_keeps_occurrences_but_reuses_analysis()
    await test_disabled_vision_never_constructs_or_calls_client()
    await test_one_image_failure_does_not_discard_other_image_results()
    await test_lease_loss_prevents_next_provider_call()
    print("document_vision_processing=passed")


if __name__ == "__main__":
    asyncio.run(main())
